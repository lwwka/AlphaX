from __future__ import annotations

import csv
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from jinja2 import Environment, FileSystemLoader, select_autoescape
from markdown import markdown
from openpyxl import Workbook
from docx import Document

from src.clients.llm_client import LLMClient
from src.models.schemas import PriceSnapshot, SentimentResult, SignalResult, TweetRecord
from src.utils.io import write_text


def build_daily_report(
    report_date: str,
    tweets: list[TweetRecord],
    sentiments: list[SentimentResult],
    signals: list[SignalResult],
    prices: dict[str, PriceSnapshot],
    settings: dict[str, Any],
    api_key: str | None,
    logger,
) -> dict[str, str]:
    context = _build_context(report_date, tweets, sentiments, signals, prices, settings, summary="")
    context["summary"] = _build_summary(context, api_key, settings, logger)

    template_env = Environment(
        loader=FileSystemLoader(Path(__file__).parent / "templates"),
        autoescape=select_autoescape(disabled_extensions=("j2",)),
        trim_blocks=True,
        lstrip_blocks=True,
    )
    template = template_env.get_template("daily_report.md.j2")
    markdown_content = template.render(**context)
    html_content = markdown(markdown_content, extensions=["tables", "fenced_code"])

    reports_dir = settings["paths"]["reports_dir"]
    filename_suffix = _build_report_filename_suffix(report_date, settings)
    md_path = f"{reports_dir}/daily_{filename_suffix}.md"
    html_path = f"{reports_dir}/daily_{filename_suffix}.html"
    csv_path = f"{reports_dir}/daily_{filename_suffix}.csv"
    excel_path = f"{reports_dir}/daily_{filename_suffix}.xlsx"
    word_path = f"{reports_dir}/daily_{filename_suffix}.docx"
    write_text(md_path, markdown_content)
    write_text(html_path, html_content)
    _write_csv_report(csv_path, context)
    _write_excel_report(excel_path, context)
    _write_word_report(word_path, context)
    return {
        "markdown": md_path,
        "html": html_path,
        "csv": csv_path,
        "excel": excel_path,
        "word": word_path,
    }


def _build_summary(
    context: dict[str, Any],
    api_key: str | None,
    settings: dict[str, Any],
    logger,
) -> str:
    if not api_key:
        return _build_fallback_summary(context)

    llm_settings = settings["llm"]
    client = LLMClient(
        base_url=llm_settings["api_base_url"],
        api_key=api_key,
        timeout_seconds=int(llm_settings.get("timeout_seconds", 45)),
        sentiment_model=llm_settings["sentiment_model"],
        report_model=llm_settings["report_model"],
        cache_dir=settings["paths"]["cache_dir"],
        provider=llm_settings["provider"],
    )
    try:
        summary = client.generate_report_summary(
            {
                "report_date": context["report_date"],
                "tweet_count": context["tweet_count"],
                "sentiment_count": context["sentiment_count"],
                "signal_count": context["signal_count"],
                "handle_counts": context["handle_counts"],
                "signals": context["signal_rows"][:10],
                "details": context["details"][:10],
                "market_snapshots": context["market_snapshots"][:12],
                "tweet_preview": context["tweet_preview"][:8],
                "coverage_note": context["coverage_note"],
            }
        ).strip()
        return summary or _build_fallback_summary(context)
    except Exception as exc:
        logger.warning("report model failed, using fallback summary: %s", exc)
        return _build_fallback_summary(context)


def _build_context(
    report_date: str,
    tweets: list[TweetRecord],
    sentiments: list[SentimentResult],
    signals: list[SignalResult],
    prices: dict[str, PriceSnapshot],
    settings: dict[str, Any],
    summary: str,
) -> dict[str, Any]:
    tweet_index = {tweet.tweet_id: tweet for tweet in tweets}
    sentiment_index = {(item.tweet_id, item.symbol): item for item in sentiments}

    details: list[dict[str, Any]] = []
    for signal in signals:
        tweet = tweet_index.get(signal.tweet_id)
        sentiment = sentiment_index.get((signal.tweet_id, signal.symbol))
        price = prices.get(signal.symbol)
        if tweet is None or sentiment is None or price is None:
            continue
        details.append(
            {
                "symbol": signal.symbol,
                "signal": signal.signal,
                "handle": signal.handle,
                "tweet_text": tweet.text,
                "tweet_time": tweet.created_at.isoformat(),
                "source_type": tweet.source_type,
                "source_weight": signal.source_weight,
                "sentiment_score": sentiment.score,
                "sentiment_label": sentiment.label,
                "rationale": sentiment.rationale,
                "signal_type": sentiment.signal_type,
                "price_close": price.close,
                "price_change_pct": round(price.change_pct, 3),
                "volume": price.volume,
                "avg_20d_volume": round(price.avg_20d_volume, 2),
                "price_source": price.source,
                "account_weight": signal.account_weight,
                "volume_factor": round(signal.volume_factor, 2),
                "price_confirmed": signal.price_confirmed,
                "explain": signal.explain,
            }
        )

    signal_rows = [
        {
            "symbol": item.symbol,
            "signal": item.signal,
            "score": round(item.final_score, 3),
            "trigger": f"@{item.handle} ({item.source_type})",
        }
        for item in signals
    ]
    handle_counts = _count_handles(tweets)
    tweet_preview = [
        {
            "handle": tweet.handle,
            "tweet_time": tweet.created_at.isoformat(),
            "text": tweet.text[:220],
        }
        for tweet in tweets[:8]
    ]
    coverage_note = _build_coverage_note(tweets, sentiments, signals)
    analysis = _build_analysis_sections(details, signal_rows, handle_counts, coverage_note)
    market_snapshots = _build_market_snapshot_rows(prices, settings)
    return {
        "report_date": report_date,
        "summary": summary,
        "signal_rows": signal_rows,
        "details": details,
        "market_snapshots": market_snapshots,
        "tweet_count": len(tweets),
        "sentiment_count": len(sentiments),
        "signal_count": len(signals),
        "handle_counts": handle_counts,
        "tweet_preview": tweet_preview,
        "coverage_note": coverage_note,
        "analysis": analysis,
        "provider": settings["llm"]["provider"],
        "sentiment_model": settings["llm"]["sentiment_model"],
        "report_model": settings["llm"]["report_model"],
    }


def _build_report_filename_suffix(report_date: str, settings: dict[str, Any]) -> str:
    timezone_name = settings.get("app", {}).get("timezone", "UTC")
    try:
        now = datetime.now(ZoneInfo(timezone_name))
    except Exception:
        now = datetime.now(timezone.utc)
    return f"{report_date}_{now.strftime('%H%M%S')}"


def _write_csv_report(path: str, context: dict[str, Any]) -> None:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    with target.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=_report_columns())
        writer.writeheader()
        for row in _report_rows(context):
            writer.writerow(row)


def _write_excel_report(path: str, context: dict[str, Any]) -> None:
    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "Summary"
    summary_sheet.append(["report_date", context["report_date"]])
    summary_sheet.append(["summary", context["summary"]])
    summary_sheet.append(["executive_takeaway", context["analysis"]["executive_takeaway"]])
    summary_sheet.append(["signal_count", len(context["signal_rows"])])
    summary_sheet.append(["detail_count", len(context["details"])])
    summary_sheet.append(["market_snapshot_count", len(context["market_snapshots"])])
    summary_sheet.append(["provider", context["provider"]])
    summary_sheet.append(["sentiment_model", context["sentiment_model"]])
    summary_sheet.append(["report_model", context["report_model"]])
    summary_sheet.append(["driver_1", context["analysis"]["drivers"][0] if context["analysis"]["drivers"] else ""])
    summary_sheet.append(["risk_1", context["analysis"]["risks"][0] if context["analysis"]["risks"] else ""])
    summary_sheet.append(["watch_1", context["analysis"]["watch_items"][0] if context["analysis"]["watch_items"] else ""])

    signals_sheet = workbook.create_sheet("Signals")
    columns = _report_columns()
    signals_sheet.append(columns)
    for row in _report_rows(context):
        signals_sheet.append([row[column] for column in columns])

    market_sheet = workbook.create_sheet("Market Snapshots")
    market_columns = _market_snapshot_columns()
    market_sheet.append(market_columns)
    for row in context["market_snapshots"]:
        market_sheet.append([row[column] for column in market_columns])

    for sheet in workbook.worksheets:
        _autosize_worksheet(sheet)

    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(target)


def _write_word_report(path: str, context: dict[str, Any]) -> None:
    document = Document()
    document.add_heading(f"每日投資信號報告 - {context['report_date']}", level=0)

    document.add_heading("摘要", level=1)
    for paragraph in _summary_paragraphs(context["summary"]):
        document.add_paragraph(paragraph)

    document.add_heading("本次覆蓋情況", level=1)
    document.add_paragraph(
        f"本次共抓取 {context['tweet_count']} 則推文，完成 {context['sentiment_count']} 筆情緒分析，產生 {context['signal_count']} 筆信號。"
    )
    document.add_paragraph(context["coverage_note"])

    if context["market_snapshots"]:
        document.add_heading("市場快照", level=1)
        market_table = document.add_table(rows=1, cols=6)
        market_table.style = "Table Grid"
        header_cells = market_table.rows[0].cells
        for index, title in enumerate(["股票", "價格", "漲跌幅", "成交量", "20 日均量", "來源"]):
            header_cells[index].text = title
        for row in context["market_snapshots"]:
            cells = market_table.add_row().cells
            cells[0].text = str(row["symbol"])
            cells[1].text = str(row["close"])
            cells[2].text = f'{row["change_pct"]}%'
            cells[3].text = str(row["volume"])
            cells[4].text = str(row["avg_20d_volume"])
            cells[5].text = str(row["source"])

    document.add_heading("一句話結論", level=1)
    document.add_paragraph(context["analysis"]["executive_takeaway"])

    document.add_heading("核心驅動", level=1)
    for line in context["analysis"]["drivers"]:
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("風險與保留", level=1)
    for line in context["analysis"]["risks"]:
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("明日追蹤", level=1)
    for line in context["analysis"]["watch_items"]:
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("信號匯總", level=1)
    summary_table = document.add_table(rows=1, cols=4)
    summary_table.style = "Table Grid"
    header_cells = summary_table.rows[0].cells
    for index, title in enumerate(["股票", "信號", "分數", "主要觸發"]):
        header_cells[index].text = title
    for row in context["signal_rows"]:
        cells = summary_table.add_row().cells
        cells[0].text = str(row["symbol"])
        cells[1].text = str(row["signal"])
        cells[2].text = str(row["score"])
        cells[3].text = str(row["trigger"])

    document.add_heading("信號詳情", level=1)
    if not context["details"]:
        document.add_paragraph("本次沒有生成可交易信號。")
    for item in context["details"]:
        document.add_heading(f"{item['symbol']} - {item['signal']}", level=2)
        document.add_paragraph(f"觸發推文：@{item['handle']} | {item['tweet_time']}")
        document.add_paragraph(f"內容來源：{item['source_type']} | 來源權重：{item['source_weight']}")
        document.add_paragraph(f"推文內容：{item['tweet_text']}")
        document.add_paragraph(
            f"情緒分數：{item['sentiment_score']}（{item['sentiment_label']}） | 信號類型：{item['signal_type']}"
        )
        document.add_paragraph(
            f"股價：{item['price_close']}（{item['price_change_pct']}%） | 成交量：{item['volume']} | 20 日均量：{item['avg_20d_volume']}"
        )
        document.add_paragraph(f"情緒說明：{item['rationale']}")
        document.add_paragraph(f"計算說明：{item['explain']}")

    document.add_heading("模型與來源", level=1)
    document.add_paragraph(f"provider: {context['provider']}")
    document.add_paragraph(f"sentiment_model: {context['sentiment_model']}")
    document.add_paragraph(f"report_model: {context['report_model']}")
    document.add_paragraph("本報告非投資建議，所有信號為算法生成，請獨立核實。")

    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def _report_columns() -> list[str]:
    return [
        "symbol",
        "signal",
        "handle",
        "tweet_time",
        "source_type",
        "source_weight",
        "tweet_text",
        "sentiment_score",
        "sentiment_label",
        "rationale",
        "signal_type",
        "price_close",
        "price_change_pct",
        "volume",
        "avg_20d_volume",
        "price_source",
        "account_weight",
        "volume_factor",
        "price_confirmed",
        "explain",
    ]


def _report_rows(context: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        {column: item.get(column, "") for column in _report_columns()}
        for item in context["details"]
    ]


def _market_snapshot_columns() -> list[str]:
    return ["symbol", "close", "change_pct", "volume", "avg_20d_volume", "source", "as_of"]


def _build_market_snapshot_rows(
    prices: dict[str, PriceSnapshot],
    settings: dict[str, Any],
) -> list[dict[str, Any]]:
    markets = settings.get("markets", {})
    preferred_order = list(markets.get("watchlist", [])) + list(markets.get("benchmarks", []))
    ordered_symbols = [symbol for symbol in preferred_order if symbol in prices]
    ordered_symbols.extend(sorted(symbol for symbol in prices if symbol not in set(ordered_symbols)))
    return [
        {
            "symbol": symbol,
            "close": round(prices[symbol].close, 4),
            "change_pct": round(prices[symbol].change_pct, 3),
            "volume": round(prices[symbol].volume, 2),
            "avg_20d_volume": round(prices[symbol].avg_20d_volume, 2),
            "source": prices[symbol].source,
            "as_of": prices[symbol].as_of.isoformat(),
        }
        for symbol in ordered_symbols
    ]


def _autosize_worksheet(sheet) -> None:
    for column_cells in sheet.columns:
        values = [str(cell.value or "") for cell in column_cells]
        width = min(max(len(value) for value in values) + 2, 80)
        column_letter = column_cells[0].column_letter
        sheet.column_dimensions[column_letter].width = width


def _count_handles(tweets: list[TweetRecord]) -> list[dict[str, Any]]:
    counts: dict[str, int] = {}
    for tweet in tweets:
        counts[tweet.handle] = counts.get(tweet.handle, 0) + 1
    return [
        {"handle": handle, "tweets": count}
        for handle, count in sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    ]


def _build_coverage_note(
    tweets: list[TweetRecord],
    sentiments: list[SentimentResult],
    signals: list[SignalResult],
) -> str:
    if not tweets:
        return "本次沒有抓到任何可用推文，建議先檢查資料來源、API 配額或追蹤帳號設定。"
    if tweets and not sentiments:
        return "本次已抓到推文，但多數內容未命中股票映射規則，所以還未進入情緒分析與信號生成。"
    if sentiments and not signals:
        return "本次已有情緒分析結果，但未達到信號門檻，屬於觀察期而非明確交易提示。"
    return "本次資料鏈路已完整跑通，以下內容可作為後續研究與人工判讀的起點。"


def _build_analysis_sections(
    details: list[dict[str, Any]],
    signal_rows: list[dict[str, Any]],
    handle_counts: list[dict[str, Any]],
    coverage_note: str,
) -> dict[str, Any]:
    drivers = _build_driver_lines(details, signal_rows, handle_counts)
    risks = _build_risk_lines(details, signal_rows, handle_counts)
    watch_items = _build_watch_items(details, signal_rows)
    return {
        "executive_takeaway": _build_executive_takeaway(details, signal_rows, handle_counts, coverage_note),
        "drivers": drivers,
        "risks": risks,
        "watch_items": watch_items,
    }


def _build_executive_takeaway(
    details: list[dict[str, Any]],
    signal_rows: list[dict[str, Any]],
    handle_counts: list[dict[str, Any]],
    coverage_note: str,
) -> str:
    if not signal_rows:
        return coverage_note

    symbol_counts = _count_values(item["symbol"] for item in signal_rows)
    top_symbol, top_count = max(symbol_counts.items(), key=lambda item: item[1])
    total_signals = len(signal_rows)
    avg_score = sum(float(item["score"]) for item in signal_rows) / total_signals
    top_handle = handle_counts[0]["handle"] if handle_counts else "unknown"
    return (
        f"今天共形成 {total_signals} 筆信號，重心明顯集中在 {top_symbol} "
        f"（{top_count}/{total_signals} 筆，平均分數 {avg_score:.2f}）。"
        f"目前主要訊號來源仍是 @{top_handle}，適合視為同一主題鏈條的延伸，而不是多個完全獨立的催化。"
    )


def _build_driver_lines(
    details: list[dict[str, Any]],
    signal_rows: list[dict[str, Any]],
    handle_counts: list[dict[str, Any]],
) -> list[str]:
    if not details:
        return ["本次未形成可交易信號，因此暫時沒有可歸納的核心驅動因素。"]

    symbol_counts = _count_values(item["symbol"] for item in signal_rows)
    top_symbol, top_count = max(symbol_counts.items(), key=lambda item: item[1])
    total_signals = len(signal_rows)
    themes = _extract_theme_counts(details)
    signal_types = _count_values(item["signal_type"] for item in details)
    top_signal_type, top_signal_type_count = max(signal_types.items(), key=lambda item: item[1])

    lines = [
        f"{top_symbol} 是本次最集中的標的，佔全部信號 {top_count}/{total_signals}，代表市場解讀高度集中，主線很清楚。",
        f"最常出現的信號類型是「{top_signal_type}」共 {top_signal_type_count} 次，表示本輪更偏向產品進展、服務擴張或技術里程碑，而不是財報或估值重估。",
    ]
    if themes:
        theme_text = "、".join(f"{theme}（{count}）" for theme, count in themes[:3])
        lines.append(f"推文題材主要集中在：{theme_text}，顯示本輪情緒並非全面擴散，而是圍繞少數敘事反覆強化。")
    if handle_counts:
        lines.append(
            "訊號來源分布為："
            + "、".join(f"@{item['handle']} {item['tweets']} 則" for item in handle_counts[:3])
            + "。若後續來源仍然單一，判讀時應把它視為單帳號敘事延伸。"
        )
    return lines[:4]


def _build_risk_lines(
    details: list[dict[str, Any]],
    signal_rows: list[dict[str, Any]],
    handle_counts: list[dict[str, Any]],
) -> list[str]:
    if not details:
        return ["本次沒有明確信號，最大風險是樣本不足，容易把零散內容過度解讀。"]

    lines: list[str] = []
    proxy_count = 0
    for item in details:
        text = item["tweet_text"].lower()
        if item["symbol"] == "TSLA" and any(keyword in text for keyword in ("spacex", "starlink", "grok", "xai")):
            proxy_count += 1
    if proxy_count:
        lines.append(
            f"{proxy_count} 筆 {signal_rows[0]['symbol']} 相關信號其實來自 SpaceX、Starlink 或 Grok 等未上市主題，"
            "屬 proxy 映射，不能直接等同於 Tesla 基本面利好。"
        )
    symbol_counts = _count_values(item["symbol"] for item in signal_rows)
    if symbol_counts:
        top_symbol, top_count = max(symbol_counts.items(), key=lambda item: item[1])
        if top_count == len(signal_rows):
            lines.append(f"所有信號都落在 {top_symbol}，集中度過高，容易放大單一敘事的偏誤。")
        elif top_count / len(signal_rows) >= 0.6:
            lines.append(f"{top_symbol} 佔比超過六成，代表組合分散度不足，報告更像主題觀察而不是全面市場掃描。")
    if handle_counts and handle_counts[0]["tweets"] == sum(item["tweets"] for item in handle_counts):
        lines.append(f"本輪內容幾乎全部來自 @{handle_counts[0]['handle']} 單一來源，消息面集中，缺少交叉驗證。")
    if any("RT @" in item["tweet_text"] for item in details):
        lines.append("仍有部份轉述式內容混入樣本，後續若要更嚴格，可再把 quoted/轉述內容降權。")
    return lines[:4] or ["目前未觀察到特別突出的風險訊號，但仍應注意模型映射與樣本集中度。"]


def _build_watch_items(details: list[dict[str, Any]], signal_rows: list[dict[str, Any]]) -> list[str]:
    if not details:
        return [
            "先觀察明天是否出現更直接的公司、產品或政策關鍵字，避免在低樣本情況下做過度判讀。",
            "若後續仍持續沒有信號，可優先擴充 entity map，而不是直接放寬門檻。",
        ]

    symbol_counts = _count_values(item["symbol"] for item in signal_rows)
    top_symbol, _ = max(symbol_counts.items(), key=lambda item: item[1])
    watch_items = [
        f"觀察 {top_symbol} 接下來是否出現更直接、非 proxy 的催化，例如產品、交付、財務或政策層面的新訊息。",
        f"觀察 {top_symbol} 在下一個交易日能否延續目前報告中的價格確認，否則這輪訊號可能只屬短線情緒反應。",
    ]
    if any(item["symbol"] == "NVDA" for item in signal_rows):
        watch_items.append("若 AI 題材持續升溫，值得再看 NVDA 是否有第二來源或更直接的需求證據，而不是只靠單條推文外溢。")
    if any("tariff" in item["tweet_text"].lower() or "china" in item["tweet_text"].lower() for item in details):
        watch_items.append("若明天開始出現更密集的關稅或中國相關內容，可再觀察 FXI、BABA 或 9988.HK 是否需要拉高追蹤權重。")
    return watch_items[:4]


def _extract_theme_counts(details: list[dict[str, Any]]) -> list[tuple[str, int]]:
    theme_keywords = {
        "Tesla 產品/保險": ("tesla", "insurance", "fsd"),
        "SpaceX 發射/里程碑": ("spacex", "falcon", "launch", "satellite"),
        "Starlink 服務擴張": ("starlink", "service", "mini"),
        "AI / Grok": ("grok", "ai", "voice-based"),
        "中國 / 關稅": ("china", "chinese", "tariff"),
        "加密貨幣": ("bitcoin", "btc", "dogecoin", "crypto"),
    }
    counts: dict[str, int] = {}
    for item in details:
        text = item["tweet_text"].lower()
        for theme, keywords in theme_keywords.items():
            if any(keyword in text for keyword in keywords):
                counts[theme] = counts.get(theme, 0) + 1
    return sorted(counts.items(), key=lambda item: (-item[1], item[0]))


def _count_values(values) -> dict[str, int]:
    counts: dict[str, int] = {}
    for value in values:
        text = str(value)
        counts[text] = counts.get(text, 0) + 1
    return counts


def _build_fallback_summary(context: dict[str, Any]) -> str:
    lines = [
        f"今天共追蹤到 {context['tweet_count']} 則推文，完成 {context['sentiment_count']} 筆情緒分析，產生 {context['signal_count']} 筆信號。",
        context["coverage_note"],
    ]
    if context["handle_counts"]:
        handle_text = "、".join(
            f"@{item['handle']} {item['tweets']} 則"
            for item in context["handle_counts"]
        )
        lines.append(f"本次推文主要來自：{handle_text}。")
    if context["signal_rows"]:
        top_signal = context["signal_rows"][0]
        lines.append(
            f"目前最值得先看的標的是 {top_signal['symbol']}，信號為 {top_signal['signal']}，分數 {top_signal['score']}，由 {top_signal['trigger']} 觸發。"
        )
        lines.append(context["analysis"]["executive_takeaway"])
    else:
        lines.append("目前未形成明確交易信號，但報告中的推文內容與覆蓋情況仍可作為後續人工研究的參考。")
    return "\n\n".join(lines)


def _summary_paragraphs(summary: str) -> list[str]:
    paragraphs = [item.strip(" -") for item in summary.split("\n\n") if item.strip()]
    return paragraphs or ["本次沒有生成摘要內容。"]
